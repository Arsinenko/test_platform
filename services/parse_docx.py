import re
import json
import docx
from docx.opc.exceptions import PackageNotFoundError
from pathlib import Path
import base64
import os
from collections import defaultdict


def extract_test_data(docx_path):
    """
    Извлекает данные тестов из документа .docx, включая вопросы, варианты ответов и правильные ответы.

    Args:
        docx_path: путь к файлу .docx

    Returns:
        словарь с данными вопросов
    """
    try:
        doc = docx.Document(docx_path)
    except PackageNotFoundError:
        print(f"Ошибка: файл {docx_path} не найден или не является допустимым документом .docx")
        return {}

    # Извлечение всего текста из документа
    # full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    full_text = extract_text_with_images(doc)

    # Начальная обработка изображений
    image_data = extract_images(doc, docx_path)

    # Разделение текста на вопросы
    return process_questions(full_text, image_data)


def extract_text_with_images(doc):
    full_text = ""
    image_index = 0
    for paragraph in doc.paragraphs:
        p_text = ""
        for run in paragraph.runs:
            # Проверяем, содержит ли run элемент изображения
            if run._element.xpath('.//w:drawing'):
                image_index += 1
                placeholder = f"[IMAGE_{image_index}]"
                p_text += placeholder
            else:
                p_text += run.text
        full_text += p_text + "\n"
    return full_text


def extract_images(doc, docx_path):
    """
    Извлекает изображения из документа .docx и сохраняет их во временную директорию.

    Args:
        doc: объект документа docx
        docx_path: путь к файлу .docx

    Returns:
        словарь соответствия между placeholder-ами изображений и путями к сохраненным изображениям
    """
    image_data = {}

    # Создаем временную директорию для изображений, если её нет
    temp_dir = Path("temp_images")
    temp_dir.mkdir(exist_ok=True)

    # Имя документа для создания уникальных имен файлов
    doc_name = Path(docx_path).stem

    # Извлечение изображений
    image_index = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_index += 1
                # Генерируем имя файла изображения
                image_filename = f"{doc_name}_image_{image_index}.png"
                image_path = temp_dir / image_filename

                # Чтение и сохранение данных изображения
                with open(image_path, 'wb') as img_file:
                    img_file.write(rel.target_part.blob)

                # Создаем placeholder для изображения, который будет использоваться в тексте
                placeholder = f"[IMAGE_{image_index}]"
                image_data[placeholder] = str(image_path)
            except Exception as e:
                print(f"Ошибка при извлечении изображения: {e}")

    return image_data


def process_questions(text, image_data):
    """
    Обрабатывает текст документа и выделяет вопросы, варианты ответов и правильные ответы.

    Args:
        text: текст документа
        image_data: словарь с данными изображений

    Returns:
        словарь с данными вопросов
    """
    # Модифицированный шаблон для поиска начала вопросов, более соответствующий формату документа
    question_pattern = r"(\d+)\.\s+Тип\s+\d+\s+№\s+\[\[\d+\]"

    # Для отладки - выведем часть текста
    print("Первые 200 символов текста:", text[:200])

    # Находим начальные позиции всех вопросов
    question_positions = [m.start() for m in re.finditer(question_pattern, text)]
    print(f"Найдено вопросов: {len(question_positions)}")

    # Если нет вопросов, пробуем альтернативный шаблон
    if not question_positions:
        alternative_pattern = r"(\d+)\.\s+Тип"
        question_positions = [m.start() for m in re.finditer(alternative_pattern, text)]
        print(f"Альтернативный поиск нашел вопросов: {len(question_positions)}")

        # Если все еще нет вопросов, пробуем еще один шаблон
        if not question_positions:
            simpler_pattern = r"(\d+)\.\s+Тип\s+\d+"
            question_positions = [m.start() for m in re.finditer(simpler_pattern, text)]
            print(f"Простой поиск нашел вопросов: {len(question_positions)}")

    # Если нет вопросов, возвращаем пустой словарь
    if not question_positions:
        print("Не удалось найти вопросы в документе. Проверьте формат документа.")
        return {}

    # Добавляем конец текста как последнюю позицию
    question_positions.append(len(text))

    # Обрабатываем каждый вопрос
    questions = []
    for i in range(len(question_positions) - 1):
        question_text = text[question_positions[i]:question_positions[i + 1]].strip()

        # Извлекаем номер вопроса
        question_number_match = re.search(r"(\d+)\.\s+Тип", question_text)
        if question_number_match:
            question_number = int(question_number_match.group(1))
        else:
            question_number = i + 1

        # Для отладки печатаем начало каждого вопроса
        print(f"Вопрос {question_number} начинается с: {question_text[:50]}...")

        # Обрабатываем текст вопроса, варианты ответов и правильный ответ
        question_data = extract_question_data(question_text, question_number, image_data)
        questions.append(question_data)

    # Формируем итоговый словарь с данными
    result = {
        "total_questions": len(questions),
        "questions": questions
    }

    return result


def extract_question_data(question_text, question_number, image_data):
    """
    Извлекает данные из текста вопроса, включая варианты ответов и правильный ответ.

    Args:
        question_text: текст вопроса
        question_number: номер вопроса
        image_data: словарь с данными изображений

    Returns:
        словарь с данными вопроса
    """
    # Базовая структура данных вопроса
    question_data = {
        "id": question_number,
        "text": "",
        "has_image": False,
        "images": [],
        "options": [],
        "correct_answer": None
    }

    # Обработка текста вопроса
    # Убираем номер вопроса и ссылку на тип
    cleaned_text = re.sub(r"\d+\.\s+Тип\s+\d+\s+№\s+\[\[\d+\]\].*?\n", "", question_text, 1)

    # Если первый шаблон не сработал, попробуем другой
    if cleaned_text == question_text:
        cleaned_text = re.sub(r"\d+\.\s+Тип\s+\d+.*?\n", "", question_text, 1)

    # Ищем правильный ответ
    answer_match = re.search(r"Ответ:\s+(.+?)(?:\.|$)", cleaned_text)
    if answer_match:
        correct_answer = answer_match.group(1).strip()
        # Удаляем строку с ответом из текста вопроса
        cleaned_text = re.sub(r"Ответ:\s+.+?(?:\.|$)", "", cleaned_text)
        question_data["correct_answer"] = correct_answer

    # Ищем варианты ответов (если они есть)
    options = []
    options_match = re.findall(r"(\d+)\)\s+(.*?)(?=\n\d+\)|$|\n\s*\n|\n*Ответ:)", cleaned_text, re.DOTALL)

    if options_match:
        for option_num, option_text in options_match:
            options.append({
                "number": int(option_num),
                "text": option_text.strip()
            })
        question_data["options"] = options

    # Проверяем наличие изображений в тексте вопроса
    for placeholder, image_path in image_data.items():
        if placeholder in cleaned_text:
            cleaned_text = cleaned_text.replace(placeholder, f"[Изображение: {image_path}]")
            question_data["has_image"] = True
            question_data["images"].append(image_path)

    # Окончательная очистка текста вопроса
    question_text_without_options = cleaned_text
    if options:
        # Удаляем варианты ответов из текста вопроса
        for option_num, option_text in options_match:
            pattern = rf"{option_num}\)\s+{re.escape(option_text.strip())}"
            question_text_without_options = re.sub(pattern, "", question_text_without_options, flags=re.DOTALL)

    # Убираем лишние пробелы и переносы строк
    question_data["text"] = re.sub(r'\n\s*\n', '\n', question_text_without_options).strip()

    return question_data


def save_to_json(data, output_path):
    """
    Сохраняет данные в файл JSON.

    Args:
        data: данные для сохранения
        output_path: путь к выходному файлу
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"Данные сохранены в {output_path}")


def parse_file(docx_path: str):
    # Путь к документу .docx (можно изменить на аргумент командной строки)

    # Путь для сохранения результата
    output_path = "data/output.json"

    # Извлечение данных
    data = extract_test_data(docx_path)

    # Сохранение результатов
    save_to_json(data, output_path)

    print(f"Обработано {data.get('total_questions', 0)} вопросов")


